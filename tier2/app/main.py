from flask import Flask, request, jsonify
import json
import sqlite3
from docx import Document
import pandas as pd
from fpdf import FPDF
import os

app = Flask(__name__)

# Database connection
DB_PATH = "/data/user_data.db"

def init_db():
    """Initialize the database."""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            email TEXT NOT NULL,
            age INTEGER NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

def save_to_db(data):
    """Save user data to the database."""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        INSERT INTO users (name, email, age) VALUES (?, ?, ?)
    ''', (data['name'], data['email'], data['age']))
    conn.commit()
    conn.close()

def generate_html(data):
    """Generate an HTML file from user data."""
    html_content = f'''
    <html>
        <body>
            <h1>User Information</h1>
            <p>Name: {data['name']}</p>
            <p>Email: {data['email']}</p>
            <p>Age: {data['age']}</p>
        </body>
    </html>
    '''
    with open("/data/user_info.html", "w") as file:
        file.write(html_content)

def generate_docx(data):
    """Generate a DOCX file from user data."""
    doc = Document()
    doc.add_heading('User Information', 0)
    doc.add_paragraph(f"Name: {data['name']}")
    doc.add_paragraph(f"Email: {data['email']}")
    doc.add_paragraph(f"Age: {data['age']}")
    doc.save("/data/user_info.docx")

def generate_pdf(data):
    """Generate a PDF file from user data."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="User Information", ln=True, align="C")
    pdf.cell(200, 10, txt=f"Name: {data['name']}", ln=True)
    pdf.cell(200, 10, txt=f"Email: {data['email']}", ln=True)
    pdf.cell(200, 10, txt=f"Age: {data['age']}", ln=True)
    pdf.output("/data/user_info.pdf")

def generate_excel(data):
    """Generate an Excel file from user data."""
    df = pd.DataFrame([data])
    df.to_excel("/data/user_info.xlsx", index=False)

@app.route('/submit', methods=['POST'])
def submit():
    """Handle form submission."""
    data = {
        'name': request.form['name'],
        'email': request.form['email'],
        'age': request.form['age']
    }

    # Save to database
    save_to_db(data)

    # Generate files
    generate_html(data)
    generate_docx(data)
    generate_pdf(data)
    generate_excel(data)

    return jsonify({"message": "Data processed successfully!"})

if __name__ == '__main__':
    init_db()
    app.run(host='0.0.0.0', port=5000)